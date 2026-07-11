from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
MANUSCRIPT = ROOT / "Manuscript_Q1_rewrite.md"
DOCX_FILES = [ROOT / "Manuscript_Q1_rewrite.docx", ROOT / "Supplementary_Information_Q1_rewrite.docx"]


def citation_audit(text: str) -> dict[str, object]:
    body, references = text.split("## References", maxsplit=1)
    citation_groups = re.findall(r"\[([0-9,\-]+)\]", body)
    cited_values: set[int] = set()
    for group in citation_groups:
        for part in group.split(","):
            if "-" in part:
                left, right = (int(value) for value in part.split("-", maxsplit=1))
                cited_values.update(range(left, right + 1))
            else:
                cited_values.add(int(part))
    cited = sorted(cited_values)
    reference_count = len(re.findall(r"^\d+\. ", references, flags=re.M))
    return {
        "reference_count": reference_count,
        "cited_numbers": cited,
        "uncited_reference_numbers": [i for i in range(1, reference_count + 1) if i not in cited],
        "out_of_range_citations": [i for i in cited if i > reference_count],
    }


def docx_audit(path: Path) -> dict[str, object]:
    with zipfile.ZipFile(path) as archive:
        corrupt_member = archive.testzip()
    document = Document(path)
    return {
        "file": path.name,
        "zip_integrity": "OK" if corrupt_member is None else f"CORRUPT: {corrupt_member}",
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "text_characters": sum(len(paragraph.text) for paragraph in document.paragraphs),
    }


def main() -> None:
    text = MANUSCRIPT.read_text(encoding="utf-8")
    audit = citation_audit(text)
    forbidden = ["external EFS validation", "MRD-occult high-risk subgroup", "mutual exclusivity"]
    print("Reference count:", audit["reference_count"])
    print("Cited numbers:", audit["cited_numbers"])
    print("Uncited references:", audit["uncited_reference_numbers"])
    print("Out-of-range citations:", audit["out_of_range_citations"])
    for phrase in forbidden:
        print(f"Forbidden-claim phrase '{phrase}':", phrase in text)
    for item in map(docx_audit, DOCX_FILES):
        print(item)


if __name__ == "__main__":
    main()
