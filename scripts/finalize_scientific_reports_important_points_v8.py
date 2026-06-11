from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from PIL import Image


ROOT = Path.cwd().resolve()
SOURCE = ROOT / "_ETV6-RUNX1_manuscript_v7_hard_fixes.docx"
OUTPUT_DIR = ROOT / "scientific_reports_submission_v8"
MAIN_OUTPUT = OUTPUT_DIR / "_ETV6-RUNX1_manuscript_v8_scientific_reports_ready.docx"
SUPPLEMENT_OUTPUT = OUTPUT_DIR / "ETV6-RUNX1_Supplementary_Information_v8.docx"
FAMILY_AUDIT = ROOT / "robustness_ml_analysis" / "tables" / "multiple_testing_audit.csv"


REFERENCES = [
    "1. Shurtleff, S. A. et al. TEL/AML1 fusion resulting from a cryptic t(12;21) is the most common genetic lesion in pediatric ALL and defines a subgroup of patients with an excellent prognosis. Leukemia 9, 1985-1989 (1995).",
    "2. Borkhardt, A. et al. Incidence and clinical relevance of TEL/AML1 fusion genes in children with acute lymphoblastic leukemia enrolled in the German and Italian multicenter therapy trials. Blood 90, 571-577 (1997).",
    "3. Loh, M. L. et al. Prospective analysis of TEL/AML1-positive patients treated on Dana-Farber Cancer Institute Consortium Protocol 95-01. Blood 107, 4508-4513 (2006).",
    "4. Borowitz, M. J. et al. Clinical significance of minimal residual disease in childhood acute lymphoblastic leukemia and its relationship to other prognostic factors: a Children's Oncology Group study. Blood 111, 5477-5485 (2008).",
    "5. Vora, A. et al. Treatment reduction for children and young adults with low-risk acute lymphoblastic leukaemia defined by minimal residual disease (UKALL 2003): a randomised controlled trial. Lancet Oncol. 14, 199-209 (2013).",
    "6. Forestier, E. et al. Outcome of ETV6/RUNX1-positive childhood acute lymphoblastic leukaemia in the NOPHO-ALL-1992 and -2000 protocols. Br. J. Haematol. 143, 268-277 (2008).",
    "7. Bokemeyer, A. et al. ETV6/RUNX1 relapse is not associated with a second hit on the ETV6- or RUNX1-allele in childhood acute lymphoblastic leukaemia. Br. J. Haematol. 167, 69-79 (2014).",
    "8. Greaves, M. F. & Wiemels, J. Origins of chromosome translocations in childhood leukaemia. Nat. Rev. Cancer 3, 639-649 (2003).",
    "9. Mori, H. et al. Chromosome translocations and covert leukemic clones are generated during normal fetal development. Proc. Natl Acad. Sci. USA 99, 8242-8247 (2002).",
    "10. Mullighan, C. G. et al. Genome-wide analysis of genetic alterations in acute lymphoblastic leukaemia. Nature 446, 758-764 (2007).",
    "11. Oksa, L. et al. Genomic determinants of therapy response in ETV6::RUNX1 leukemia. Leukemia 39, 2125-2139 (2025).",
    "12. Benjamini, Y. & Hochberg, Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J. R. Stat. Soc. B 57, 289-300 (1995).",
    "13. Benjamini, Y. & Yekutieli, D. The control of the false discovery rate in multiple testing under dependency. Ann. Stat. 29, 1165-1188 (2001).",
    "14. Holm, S. A simple sequentially rejective multiple test procedure. Scand. J. Stat. 6, 65-70 (1979).",
    "15. Cox, D. R. Regression models and life-tables. J. R. Stat. Soc. B 34, 187-220 (1972).",
    "16. Varma, S. & Simon, R. Bias in error estimation when using cross-validation for model selection. BMC Bioinformatics 7, 91 (2006).",
    "17. Hubert, L. & Arabie, P. Comparing partitions. J. Classif. 2, 193-218 (1985).",
    "18. Cerami, E. et al. The cBio cancer genomics portal: an open platform for exploring multidimensional cancer genomics data. Cancer Discov. 2, 401-404 (2012).",
    "19. Gao, J. et al. Integrative analysis of complex cancer genomics and clinical profiles using the cBioPortal. Sci. Signal. 6, pl1 (2013).",
    "20. Dempster, J. M. et al. Chronos: a cell population dynamics model of CRISPR experiments that improves inference of gene fitness effects. Genome Biol. 22, 343 (2021).",
    "21. Davidson-Pilon, C. lifelines: survival analysis in Python. J. Open Source Softw. 4, 1317 (2019).",
]


FAMILY_DESCRIPTIONS = {
    "target_karyotype_event": ("TARGET karyotype", "Observed EFS event", "Adverse-outcome screen"),
    "target_karyotype_poor5y": ("TARGET karyotype", "Poor 5-year EFS", "Adverse-outcome screen"),
    "target_karyotype_univ_cox": ("TARGET karyotype", "Univariable Cox EFS", "Adverse-outcome screen"),
    "target_karyotype_adjusted_cox": ("TARGET karyotype", "Adjusted Cox EFS", "Adverse-outcome screen"),
    "target_cna_event": ("TARGET baseline CNA", "Observed EFS event", "Adverse-outcome screen"),
    "target_cna_poor5y": ("TARGET baseline CNA", "Poor 5-year EFS", "Adverse-outcome screen"),
    "target_cna_candidate_univ_cox": ("TARGET baseline CNA candidates", "Univariable Cox EFS", "Adverse-outcome screen"),
    "target_cna_candidate_adjusted_cox": ("TARGET baseline CNA candidates", "Adjusted Cox EFS", "Adverse-outcome screen"),
    "target_rna_time_to_relapse": ("TARGET relapse-only baseline RNA", "Time to relapse correlation", "Adverse-outcome exploratory screen"),
    "target_rna_early_vs_late_relapse": ("TARGET relapse-only baseline RNA", "Early versus late relapse", "Adverse-outcome exploratory screen"),
    "target_rna_paired_relapse_change": ("TARGET paired RNA", "Baseline-to-relapse change", "Adverse-outcome exploratory screen"),
    "target_subtype_expression_previous": ("TARGET legacy subtype RNA", "ETV6::RUNX1-positive versus negative", "Legacy subtype sensitivity analysis"),
    "depmap_26q1_pooled_t": ("DepMap 26Q1 CRISPR", "Pooled t-test: del6q versus no del6q", "Dependency sensitivity analysis"),
    "depmap_26q1_welch_t": ("DepMap 26Q1 CRISPR", "Welch t-test: del6q versus no del6q", "Dependency sensitivity analysis"),
    "depmap_26q1_exact_permutation": ("DepMap 26Q1 CRISPR", "Exact label permutation", "Dependency sensitivity analysis"),
    "single_center_univ_cox": ("Historical single-center cohort", "Univariable Cox EFS", "Adverse-outcome exploratory screen"),
    "target_primary_only_subtype_expression": ("TARGET primary-only RNA", "ETV6::RUNX1-positive versus negative", "Subtype-biology analysis; not prognosis"),
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Finalize Scientific Reports v8 manuscript and supplement.")
    parser.add_argument("--workspace-root", type=Path, default=ROOT)
    parser.add_argument("--source-manuscript", type=Path)
    return parser.parse_args()


def configure(args: argparse.Namespace) -> None:
    global ROOT, SOURCE, OUTPUT_DIR, MAIN_OUTPUT, SUPPLEMENT_OUTPUT, FAMILY_AUDIT
    ROOT = args.workspace_root.resolve()
    SOURCE = args.source_manuscript.resolve() if args.source_manuscript else ROOT / "_ETV6-RUNX1_manuscript_v7_hard_fixes.docx"
    OUTPUT_DIR = ROOT / "scientific_reports_submission_v8"
    MAIN_OUTPUT = OUTPUT_DIR / "_ETV6-RUNX1_manuscript_v8_scientific_reports_ready.docx"
    SUPPLEMENT_OUTPUT = OUTPUT_DIR / "ETV6-RUNX1_Supplementary_Information_v8.docx"
    FAMILY_AUDIT = ROOT / "robustness_ml_analysis" / "tables" / "multiple_testing_audit.csv"


def find_paragraph(document: Document, prefix: str, last: bool = False) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if not matches:
        raise ValueError(f"Paragraph not found: {prefix}")
    return matches[-1] if last else matches[0]


def replace_all(document: Document, prefix: str, text: str) -> None:
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if not matches:
        raise ValueError(f"Paragraph not found: {prefix}")
    for paragraph in matches:
        paragraph.text = text


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def remove_table(table) -> None:
    parent = table._element.getparent()
    if parent is not None:
        parent.remove(table._element)


def paragraph_index(document: Document, target: Paragraph) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError("Paragraph is not in document")


def insert_paragraph_before(anchor: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def extract_image(document: Document, paragraph_prefix: str, destination: Path, last: bool = False) -> None:
    paragraph = find_paragraph(document, paragraph_prefix, last=last)
    previous = paragraph._p.getprevious()
    while previous is not None and not previous.xpath(".//a:blip"):
        previous = previous.getprevious()
    if previous is None:
        raise RuntimeError(f"Image not found before caption: {paragraph_prefix}")
    blip = previous.xpath(".//a:blip")[0]
    rel_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    destination.write_bytes(document.part.related_parts[rel_id].blob)


def optimize_for_embedding(source: Path, destination: Path, max_width: int = 2400) -> None:
    with Image.open(source) as image:
        image.load()
        if image.width > max_width:
            height = round(image.height * max_width / image.width)
            image = image.resize((max_width, height), Image.Resampling.LANCZOS)
        image.save(destination, format="PNG", optimize=True)


def add_table(document: Document, headers: list[str], rows: list[list[str]], font_size: int | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    return table


def table_to_rows(table) -> tuple[list[str], list[list[str]]]:
    headers = [cell.text for cell in table.rows[0].cells]
    rows = [[cell.text for cell in row.cells] for row in table.rows[1:]]
    return headers, rows


def count_words(text: str) -> int:
    return len(re.findall(r"\b[\w:+-]+\b", text))


def section_word_count(document: Document, start: str, end: str) -> int:
    active = False
    total = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == start:
            active = True
            continue
        if text == end:
            break
        if active and not paragraph.style.name.startswith("Heading"):
            total += count_words(text)
    return total


def build_supplement(source: Document, family_audit: pd.DataFrame, image_map: dict[str, Path]) -> None:
    supplement = Document()
    supplement.add_heading("Supplementary Information", level=1)
    supplement.add_paragraph(
        "Multiplicity-controlled reassessment of candidate adverse-outcome biomarkers in "
        "ETV6::RUNX1-positive pediatric B-ALL"
    )
    supplement.add_paragraph(
        "Supplementary items are exploratory or transparency records unless explicitly stated otherwise."
    )

    supplement.add_heading("Supplementary Tables", level=1)
    supplement.add_heading("Supplementary Table S1. Multiplicity-controlled and machine-learning reliability assessment", level=2)
    headers, rows = table_to_rows(source.tables[4])
    add_table(supplement, headers, rows)
    supplement.add_paragraph(
        "ARI: adjusted Rand index; AUC: area under the receiver-operating-characteristic curve; "
        "BH: Benjamini-Hochberg; BY: Benjamini-Yekutieli; CV: cross-validation."
    )

    supplement.add_heading("Supplementary Table S2. Complete multiple-testing family register", level=2)
    family_rows = []
    for row in family_audit.itertuples(index=False):
        data_layer, question, role = FAMILY_DESCRIPTIONS[row.family]
        family_rows.append(
            [
                row.family,
                data_layer,
                question,
                str(row.tests),
                "BH-FDR primary; BY-FDR and Holm-FWER sensitivity",
                role,
                str(row.bh_significant_0_05),
            ]
        )
    add_table(
        supplement,
        ["Family ID", "Data layer", "Endpoint / question", "Tests", "Correction", "Role", "BH discoveries"],
        family_rows,
    )
    supplement.add_paragraph(
        "Families were defined operationally by distinct data layer, endpoint, and model specification "
        "before applying correction. They were not preregistered and all inferential findings remain exploratory."
    )

    captions = {
        "S1": find_paragraph(source, "Supplementary Figure S1.").text,
        "S2": find_paragraph(source, "Supplementary Figure S2.").text,
        "S3": find_paragraph(source, "Supplementary Figure S3.").text,
        "S4": find_paragraph(source, "Supplementary Figure S4.").text,
        "S5": find_paragraph(source, "Supplementary Figure S5.").text,
        "S6": find_paragraph(source, "Supplementary Figure S6.").text,
        "S7": find_paragraph(source, "Figure 4. Single-center evidence boundaries").text.replace("Figure 4.", "Supplementary Figure S7.", 1),
        "S8": find_paragraph(source, "Figure 6. Cross-validated supervised-learning").text.replace(
            "Figure 6. Cross-validated supervised-learning results.",
            "Supplementary Figure S8. Internal cross-validation sensitivity results.",
            1,
        ),
    }
    captions["S7"] = captions["S7"].replace("(A)", "(a)").replace("(B)", "(b)").replace("(C)", "(c)").replace("(D)", "(d)")
    captions["S8"] = captions["S8"].replace("(A)", "(a)").replace("(B)", "(b)")

    supplement.add_heading("Supplementary Figures", level=1)
    embed_dir = OUTPUT_DIR / "_embed"
    embed_dir.mkdir(exist_ok=True)
    for key in [f"S{i}" for i in range(1, 9)]:
        optimized = embed_dir / f"Supplementary_Figure_{key}.png"
        optimize_for_embedding(image_map[key], optimized)
        paragraph = supplement.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(optimized), width=Inches(6.5))
        supplement.add_paragraph(captions[key])
    supplement.save(SUPPLEMENT_OUTPUT)
    for path in embed_dir.iterdir():
        if path.is_file():
            path.unlink()
    embed_dir.rmdir()


def main() -> None:
    configure(parse_args())
    if not SOURCE.exists() or not FAMILY_AUDIT.exists():
        raise FileNotFoundError(SOURCE if not SOURCE.exists() else FAMILY_AUDIT)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    main_figures = OUTPUT_DIR / "main_figures"
    supplementary_figures = OUTPUT_DIR / "supplementary_figures"
    main_figures.mkdir(exist_ok=True)
    supplementary_figures.mkdir(exist_ok=True)

    source = Document(SOURCE)
    image_map = {
        "Figure_1": main_figures / "Figure_1.png",
        "Figure_2": main_figures / "Figure_2.png",
        "Figure_3": main_figures / "Figure_3.png",
        "Figure_4": main_figures / "Figure_4.png",
        "S1": supplementary_figures / "Supplementary_Figure_S1.png",
        "S2": supplementary_figures / "Supplementary_Figure_S2.png",
        "S3": supplementary_figures / "Supplementary_Figure_S3.png",
        "S4": supplementary_figures / "Supplementary_Figure_S4.png",
        "S5": supplementary_figures / "Supplementary_Figure_S5.png",
        "S6": supplementary_figures / "Supplementary_Figure_S6.png",
        "S7": supplementary_figures / "Supplementary_Figure_S7.png",
        "S8": supplementary_figures / "Supplementary_Figure_S8.png",
    }
    extract_image(source, "Figure 1.", image_map["Figure_1"])
    extract_image(source, "Figure 2.", image_map["Figure_2"])
    extract_image(source, "Figure 3.", image_map["Figure_3"])
    extract_image(source, "Figure 5.", image_map["Figure_4"])
    for index in range(1, 7):
        extract_image(source, f"Supplementary Figure S{index}.", image_map[f"S{index}"])
    extract_image(source, "Figure 4. Single-center evidence boundaries", image_map["S7"])
    extract_image(source, "Figure 6. Cross-validated supervised-learning", image_map["S8"])

    family_audit = pd.read_csv(FAMILY_AUDIT)
    build_supplement(source, family_audit, image_map)

    document = Document(SOURCE)
    document.paragraphs[0].text = (
        "Multiplicity-controlled reassessment of candidate adverse-outcome biomarkers in "
        "ETV6::RUNX1-positive pediatric B-ALL"
    )
    document.paragraphs[1].text = "Running title: Reassessment of ETV6::RUNX1 biomarkers"
    document.paragraphs[3].text = "Figures/Tables: 4 main figures + 4 main tables"

    replace_all(
        document,
        "ETV6::RUNX1-positive pediatric B-cell acute lymphoblastic leukemia",
        "ETV6::RUNX1-positive pediatric B-cell acute lymphoblastic leukemia usually has a favorable outcome, but a subset relapses. We reassessed candidate adverse-outcome biomarkers using public TARGET, NOPHO, DepMap, and cBioPortal data and a historical single-center ETV6::RUNX1-positive outcome cohort. Reliability analyses included Benjamini-Hochberg, Benjamini-Yekutieli, and Holm correction, unsupervised stability analysis, and nested cross-validation. No karyotype, copy-number, prognostic RNA, paired RNA, DepMap, or single-center candidate survived Benjamini-Hochberg correction. Chromosome 16 gain (+16) was the leading low-frequency candidate (hazard ratio=4.20; P=0.00879; FDR=0.308), but adding +16 to a three-variable clinical-threshold Cox model yielded an unstable mean outer-fold C-index increment of 0.024 and median increment of 0.000. Karyotype clusters were weak, unstable, and unrelated to event-free survival. In primary-diagnosis RNA data (11 ETV6::RUNX1-positive and 99 negative patients), SLC7A11 was significant by Benjamini-Hochberg but not Benjamini-Yekutieli or Holm correction. Nested classification identified an internally robust fusion-subtype signal, not an outcome-prediction signal. Current data therefore do not establish a validated adverse-outcome biomarker or therapeutic mechanism; +16 warrants prespecified multicenter validation.",
    )
    replace_all(
        document,
        "Keywords:",
        "Keywords: ETV6::RUNX1, pediatric B-ALL, multiple testing, nested cross-validation, chromosome 16 gain, prognostic biomarker",
    )
    replace_all(
        document,
        "In this study, we reanalyzed",
        "In this study, we reanalyzed public TARGET, NOPHO, DepMap, and cBioPortal resources and assessed relevant candidates in a de-identified historical single-center ETV6::RUNX1-positive outcome cohort to: (1) screen recurrent karyotype features for association with EFS; (2) test the robustness and external assessability of del6q, +16, and a previously proposed multi-hit score; (3) screen baseline CNA and RNA data for candidate adverse-outcome biomarkers while controlling false discovery; (4) quantify unsupervised structure and internal out-of-sample discrimination using stability analysis and nested cross-validation; and (5) distinguish internally robust subtype observations from hypotheses requiring multicenter clinical and experimental validation.",
    )
    replace_all(
        document,
        "Study cohorts and data provenance.",
        "Study cohorts and data provenance. TARGET clinical, karyotype, CNA, and RNA-seq Z-score data were obtained from the cBioPortal DataHub study all_phase2_target_2018_pub at commit 78e9db506f1b0f029840a71712cb07dce1c51daf.[18,19] The public clinical table contained 222 ETV6::RUNX1-positive patients; 219 had valid EFS and 141 had both documented karyotype and valid EFS. Public NOPHO CNVkit-derived data comprised 262 .cns files associated with Oksa et al.[11] DepMap Public 25Q2 and 26Q1 releases were used for cross-version dependency sensitivity analyses because the manuscript-specified 24Q2 bulk matrix could not be retrieved from the public download endpoint. Sample-level derivation tables, input checksums, software versions, and analysis scripts are available in the public reproducibility repository.",
    )
    replace_all(
        document,
        "Clinical outcome analysis.",
        "Clinical outcome analysis. EFS was defined from diagnosis to relapse, death, or second malignant neoplasm; other observations were censored at last follow-up. Kaplan-Meier estimates, log-rank tests, and Cox proportional-hazards models were used.[15] The primary karyotype screen used the 141 patients with documented karyotype and valid EFS. Adjusted candidate models included Day-29 MRD, WBC, and age. The previously proposed three-factor score assigned one point each for del6q, MRD>=0.01%, and complex karyotype; it was rerun as a sensitivity analysis rather than treated as a validated model. Leave-one-positive-patient-out analyses assessed the stability of the low-frequency +16 signal.",
    )
    replace_all(
        document,
        "DepMap dependency sensitivity analysis.",
        "DepMap dependency sensitivity analysis. CRISPR Chronos gene-effect matrices from DepMap Public 25Q2 and 26Q1 were analyzed for eight prespecified B-ALL models.[20] Only five models (REH, NALM6, SEM, RS4;11, and 697) were available in both matrices. Pooled and Welch t tests, exact label permutation, BH-FDR correction, and leave-one-out analyses were performed. Because model availability and copy-number classification were unstable across releases, results were interpreted as sensitivity analyses rather than evidence of del6q-specific dependency.",
    )
    replace_all(
        document,
        "Statistics.",
        "Statistics. All tests were two-sided with alpha=0.05. Multiple-testing families were defined operationally by distinct data layer, endpoint, and model specification before applying correction; they were not preregistered and remain exploratory. The complete family register, including the number of tests and inferential role, is provided in Supplementary Table S2. Benjamini-Hochberg FDR was the primary discovery control, while Benjamini-Yekutieli FDR and Holm family-wise error correction were conservative sensitivity analyses.[12-14] Nominal P values describe candidate ranking and are not confirmatory when adjusted results are non-significant. Effect estimates are reported with 95% confidence intervals where available. Analyses used Python 3.11 with scipy, lifelines, statsmodels, pandas, numpy, and matplotlib.[21]",
    )
    replace_all(
        document,
        "Machine-learning reliability analysis.",
        "Machine-learning reliability analysis. Unsupervised karyotype analysis used Jaccard distance and average-linkage hierarchical clustering after prevalence filtering and removal of duplicate binary feature patterns. Candidate cluster numbers k=2-6 were compared by silhouette, outcome association was evaluated by label permutation and log-rank testing, and stability was quantified by adjusted Rand index across 300 repeated 80% feature-subsampling analyses.[17] Prognostic supervised analysis was restricted to the 141-patient TARGET karyotype cohort with 24 EFS events. Three-variable clinical-threshold, clinical-plus-+16, and clinical-plus-all-eligible-karyotype Cox models were evaluated using repeated nested five-fold cross-validation; feature filtering, imputation, and ridge-penalty selection occurred within training folds.[16] Concordance was calculated separately in each outer test fold before aggregation. Primary-only RNA subtype classification used a nearest-centroid classifier with nested feature-number selection; the complete pipeline was evaluated using 200 label permutations, and AUC was calculated separately in each outer test fold. Results are reported as internal exploratory discrimination estimates, not as validated clinical prediction models; means, medians, and fold ranges are shown because sparse events precluded precise performance intervals.",
    )

    replacements = {
        "This reproducible cohort": "This publicly derivable cohort",
        "in the reproducible cohort": "in the publicly derivable cohort",
        "reproducible 141-patient karyotype cohort": "publicly derivable 141-patient karyotype cohort",
        "This reproducible public-data reanalysis": "This multiplicity-controlled public-data reanalysis",
        "SLC7A11 upregulation remains a reproducible but correction-sensitive": "SLC7A11 upregulation is internally robust within TARGET but correction-sensitive",
        "Cross-version DepMap analysis does not support reproducible del6q-specific dependency": "Cross-version DepMap analysis does not support a del6q-specific dependency",
        "provide reproducible FDR-significant": "provide cross-version FDR-significant",
        "incompletely reproducible model set": "model set that was not consistently available across releases",
        "Reproducible public TARGET cohort": "Publicly derivable TARGET cohort",
        "Reproducible public-data result": "Public-data result",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text
        for old, new in replacements.items():
            text = text.replace(old, new)
        paragraph.text = text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for old, new in replacements.items():
                    text = text.replace(old, new)
                cell.text = text

    replace_all(
        document,
        "D19 MRD was available",
        "D19 MRD was available for 69 patients. The >=0.01% group contained 41 patients and five events, whereas the <0.01% group contained 28 patients and one event. The association was directional but imprecise (univariable HR=2.95, 95%CI 0.34-25.28, P=0.324). +16 was absent and del6q was observed in only one historical outcome patient without an event. Thus, the historical single-center cohort neither validates nor refutes the low-frequency +16 candidate and does not support del6q as an established adverse-outcome marker (Supplementary Figure S7; Table 4).",
    )
    replace_all(
        document,
        "For transparency, the original NOPHO breakpoint/MCR diagram",
        "For transparency, the legacy TARGET del6q/multi-hit analysis, original NOPHO breakpoint/MCR diagram, and DepMap nominal dependency plot are retained as Supplementary Figure S1, Supplementary Figure S2, and Supplementary Figure S3, respectively. The legacy TARGET panel documents a non-reconstructable earlier cohort definition and is not used for primary inference. The NOPHO diagram is a candidate-gene schematic rather than validated breakpoint mapping. The DepMap plot displays nominal results from a model set that was not consistently available across releases and must not be interpreted as evidence of significant selective dependency.",
    )
    replace_all(
        document,
        "The single-center analysis provides an independent clinical context",
        "The single-center analysis provides an independent clinical context but not a definitive external validation cohort. Its favorable long-term outcome estimates were consistent with the established prognosis of ETV6::RUNX1-positive disease. However, +16 was absent and del6q occurred only once in the outcome cohort, illustrating the practical limitation of validating rare cytogenetic candidates in a single institution. Absence of +16 does not refute the public-data signal; it indicates that a substantially larger multicenter cohort is required. A separate contemporary B-ALL registry subset lacked ETV6::RUNX1 ascertainment and is therefore reported only as descriptive phenotype and early-response context, not as biomarker validation (Supplementary Figure S6).",
    )
    replace_all(
        document,
        "SLC7A11 upregulation is internally robust within TARGET but correction-sensitive",
        "SLC7A11 upregulation is an internally robust but correction-sensitive",
    )
    replace_all(
        document,
        "No adverse-outcome screening family yielded",
        "No adverse-outcome screening family yielded a BH-FDR-significant candidate. This included the TARGET karyotype, genome-wide CNA, time-to-relapse RNA, paired relapse RNA, DepMap dependency, and single-center Cox families; BY and Holm sensitivity analyses were at least as conservative. The nominal single-center complex-karyotype association had BH-FDR=0.224, and the +16 karyotype association had BH-FDR=0.308. The complete operational family definitions and test counts are reported in Supplementary Table S2.",
    )
    replace_all(
        document,
        "Unsupervised analysis of 36 non-duplicate",
        "Unsupervised analysis of 36 non-duplicate recurrent karyotype features selected four clusters by silhouette, but separation was weak (silhouette=0.331), feature-subsampling stability was low (median adjusted Rand index=0.114), and clusters were unrelated to EFS events (permutation P=1.000; log-rank P=0.942). Repeated nested cross-validation showed limited three-variable clinical-threshold discrimination (mean outer-fold C-index=0.566). Adding +16 increased the mean outer-fold C-index to 0.590, but the paired change had mean +0.024, median +0.000, and outer-fold range -0.292 to +0.171. Adding all eligible karyotype features yielded mean outer-fold C-index=0.521 and did not improve internal discrimination (Figure 4; Supplementary Table S1).",
    )
    replace_all(
        document,
        "Primary-only RNA did not show",
        "Primary-only RNA did not show significant global unsupervised separation by fusion status (label-permutation P=0.241), but nested supervised feature selection identified a sparse fusion-subtype expression signal (mean outer-fold AUC=0.993; complete nested-pipeline permutation P=0.00498, 200 permutations). Primary sample-type distributions were similar by fusion status (Fisher P=0.725). This internally robust TARGET subtype signal has not been independently replicated and cannot be interpreted as adverse-outcome prediction (Supplementary Figure S8).",
    )
    replace_all(
        document,
        "Figure 1. Public cohort availability",
        "Figure 1. Public cohort availability and karyotype biomarker landscape. (a) Patient availability across clinical, karyotype, baseline CNA, and RNA-seq analyses. (b) Univariable Cox hazard ratios for leading unique karyotype signatures; red denotes the exploratory +16 candidate. (c) EFS-event proportions with Wilson 95% confidence intervals. (d) Univariable and clinically adjusted hazard ratios. No screened karyotype feature remained significant after BH-FDR correction.",
    )
    replace_all(
        document,
        "Figure 2. Survival and robustness",
        "Figure 2. Survival and robustness assessment of the exploratory +16 signal. (a) Kaplan-Meier EFS curves stratified by +16. (b) Kaplan-Meier EFS curves stratified by del6q as the prespecified reference hypothesis. (c) Patient-level follow-up and events among the seven +16 cases. (d) Leave-one-+16-patient-out adjusted Cox estimates. The risk direction was retained, but +16 remained low-frequency and was not FDR significant.",
    )
    replace_all(
        document,
        "Figure 5. Multiplicity-controlled",
        "Figure 4. Multiplicity-controlled reliability assessment. (a) Minimum nominal P values and BH-FDR values across analysis families. (b) Principal-component representation of unsupervised karyotype structure; clusters were weak, unstable, and unrelated to outcome. (c) Primary-only RNA principal components did not show significant global unsupervised separation by fusion status. (d) Internal nested-CV survival discrimination for the three-variable clinical-threshold, clinical-plus-+16, and clinical-plus-all-karyotype models.",
    )
    replace_all(
        document,
        "Figure 4. Single-center evidence boundaries",
        "Supplementary Figure S7. Single-center evidence boundaries for candidate adverse-outcome biomarkers.",
    )
    replace_all(
        document,
        "Figure 6. Cross-validated supervised-learning",
        "Supplementary Figure S8. Internal cross-validation sensitivity results.",
    )
    replace_all(
        document,
        "Delta Z is the difference",
        "Delta Z is the difference in mean public cBioPortal RPKM Z-score using one primary-diagnosis sample per patient. BH-FDR is genome-wide across 23,590 unique expression genes. BY-FDR and Holm correction are reported in Supplementary Table S1 and the public analysis outputs. These results describe subtype expression and do not validate prognosis or a del6q mechanism.",
    )

    # Replace the reference list with Nature-style numbered references.
    references_heading = find_paragraph(document, "References")
    figure_legends_heading = find_paragraph(document, "Figure Legends")
    for paragraph in list(document.paragraphs):
        if paragraph._p.getprevious() is references_heading._p:
            pass
    ref_start = paragraph_index(document, references_heading)
    fig_start = paragraph_index(document, figure_legends_heading)
    for paragraph in list(document.paragraphs[ref_start + 1 : fig_start]):
        remove_paragraph(paragraph)
    for reference in REFERENCES:
        insert_paragraph_before(figure_legends_heading, reference)

    # Remove all embedded image paragraphs and body-level captions.
    figure_legends_heading = find_paragraph(document, "Figure Legends")
    for paragraph in list(document.paragraphs):
        if paragraph._p.xpath(".//w:drawing"):
            remove_paragraph(paragraph)
    for relationship_id, relationship in list(document.part.rels.items()):
        if relationship.reltype.endswith("/image"):
            document.part.drop_rel(relationship_id)
    figure_legends_index = paragraph_index(document, figure_legends_heading)
    for paragraph in list(document.paragraphs[:figure_legends_index]):
        if paragraph.text.strip().startswith(("Figure ", "Supplementary Figure ")):
            remove_paragraph(paragraph)

    # Keep only main figure legends 1-4.
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text.startswith(("Supplementary Figure S7.", "Supplementary Figure S8.")):
            remove_paragraph(paragraph)

    # Remove main Table 5; it is now Supplementary Table S1.
    remove_table(document.tables[4])
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip().startswith("Table 5.") or paragraph.text.strip().startswith("ARI: adjusted Rand index"):
            remove_paragraph(paragraph)

    # Remove the embedded supplementary section; it is now a separate file.
    supplementary_heading = find_paragraph(document, "Supplementary Figures Retained for Transparency")
    supplementary_index = paragraph_index(document, supplementary_heading)
    for paragraph in list(document.paragraphs[supplementary_index:]):
        remove_paragraph(paragraph)

    abstract_words = count_words(find_paragraph(document, "ETV6::RUNX1-positive pediatric B-cell acute lymphoblastic leukemia").text)
    methods_words = section_word_count(document, "Methods", "Results")
    main_words = (
        section_word_count(document, "Introduction", "Methods")
        + section_word_count(document, "Results", "Discussion")
        + section_word_count(document, "Discussion", "Acknowledgments")
    )
    document.paragraphs[2].text = (
        f"Word counts: Abstract: {abstract_words}; Main text: {main_words}; Methods: {methods_words}"
    )
    document.save(MAIN_OUTPUT)

    manifest = f"""# Scientific Reports v8 submission manifest

## Files

- Main manuscript: `{MAIN_OUTPUT.name}`
- Supplementary information: `{SUPPLEMENT_OUTPUT.name}`
- Main figures: `main_figures/Figure_1.png` through `Figure_4.png`
- Supplementary figures: `supplementary_figures/Supplementary_Figure_S1.png` through `Supplementary_Figure_S8.png`

## Main display items

- 4 main figures
- 4 main tables
- Combined main display items: 8

## Supplementary display items

- 8 supplementary figures
- 2 supplementary tables

## Word counts

- Abstract: {abstract_words}
- Main text: {main_words}
- Methods: {methods_words}

## Remaining author-supplied blockers

- Insert ethics committee name, approval number, and consent/waiver statement.
- Complete authors, affiliations, author contributions, funding, acknowledgments, and correspondence email.
"""
    (OUTPUT_DIR / "SUBMISSION_MANIFEST.md").write_text(manifest, encoding="utf-8")
    print(MAIN_OUTPUT)
    print(SUPPLEMENT_OUTPUT)


if __name__ == "__main__":
    main()
