from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.parts.image import ImagePart
from docx.shared import Inches


ROOT = Path.cwd().resolve()
SOURCE = ROOT / "_ETV6-RUNX1_manuscript_v5_integrated_single_center.docx"
OUTPUT = ROOT / "_ETV6-RUNX1_manuscript_v7_hard_fixes.docx"
ANALYSIS = ROOT / "robustness_ml_analysis"
FIGURE_5 = ANALYSIS / "figures" / "Figure_R1_robustness_ml.png"
FIGURE_6 = ANALYSIS / "figures" / "Figure_R2_cross_validation.png"
EXPRESSION = ANALYSIS / "tables" / "primary_only_expression_statistics.csv"
SUMMARY = ANALYSIS / "analysis_summary.json"
OLD_FIGURE_3 = (
    ROOT
    / "single_center_manuscript_analysis"
    / "figures"
    / "Figure_SC1_single_center_outcomes.png"
)
NEW_FIGURE_3 = (
    ROOT
    / "single_center_manuscript_analysis"
    / "figures"
    / "Figure_SC1_historical_only_outcomes.png"
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the Scientific Reports v7 hard-fix manuscript.")
    parser.add_argument("--workspace-root", type=Path, default=ROOT)
    parser.add_argument("--source-manuscript", type=Path)
    parser.add_argument("--output-manuscript", type=Path)
    return parser.parse_args()


def configure_paths(args: argparse.Namespace) -> None:
    global ROOT, SOURCE, OUTPUT, ANALYSIS, FIGURE_5, FIGURE_6, EXPRESSION, SUMMARY
    global OLD_FIGURE_3, NEW_FIGURE_3
    ROOT = args.workspace_root.resolve()
    SOURCE = (
        args.source_manuscript.resolve()
        if args.source_manuscript
        else ROOT / "_ETV6-RUNX1_manuscript_v5_integrated_single_center.docx"
    )
    OUTPUT = (
        args.output_manuscript.resolve()
        if args.output_manuscript
        else ROOT / "_ETV6-RUNX1_manuscript_v7_hard_fixes.docx"
    )
    ANALYSIS = ROOT / "robustness_ml_analysis"
    FIGURE_5 = ANALYSIS / "figures" / "Figure_R1_robustness_ml.png"
    FIGURE_6 = ANALYSIS / "figures" / "Figure_R2_cross_validation.png"
    EXPRESSION = ANALYSIS / "tables" / "primary_only_expression_statistics.csv"
    SUMMARY = ANALYSIS / "analysis_summary.json"
    OLD_FIGURE_3 = (
        ROOT
        / "single_center_manuscript_analysis"
        / "figures"
        / "Figure_SC1_single_center_outcomes.png"
    )
    NEW_FIGURE_3 = (
        ROOT
        / "single_center_manuscript_analysis"
        / "figures"
        / "Figure_SC1_historical_only_outcomes.png"
    )


def replace_embedded_image(document: Document, old_image: Path, new_image: Path) -> None:
    old_hash = hashlib.sha256(old_image.read_bytes()).hexdigest()
    replacements = 0
    for part in document.part.package.parts:
        if isinstance(part, ImagePart):
            if hashlib.sha256(part.blob).hexdigest() == old_hash:
                part._blob = new_image.read_bytes()
                replacements += 1
    if replacements != 1:
        raise RuntimeError(f"Expected to replace one embedded Figure 3 image, replaced {replacements}")


def remove_table_row(table, row) -> None:
    table._tbl.remove(row._tr)


def find_paragraph(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def replace_paragraph(document: Document, prefix: str, text: str) -> None:
    paragraph = find_paragraph(document, prefix)
    paragraph.text = text


def replace_all_paragraphs(document: Document, prefix: str, text: str) -> None:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(prefix)
    ]
    if not matches:
        raise ValueError(f"Paragraph not found: {prefix}")
    for paragraph in matches:
        paragraph.text = text


def insert_picture_before(document: Document, anchor, image: Path, width: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    paragraph.add_run().add_picture(str(image), width=Inches(width))
    anchor._p.addprevious(paragraph._p)


def move_before(anchor, element) -> None:
    anchor._p.addprevious(element)


def fmt_p(value: float) -> str:
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.3f}"


def update_expression_table(document: Document) -> None:
    results = pd.read_csv(EXPRESSION).set_index("gene")
    table = document.tables[2]
    table.cell(0, 5).text = "Genome-wide BH-FDR"
    for row in table.rows[1:]:
        gene = row.cells[0].text.strip()
        result = results.loc[gene]
        values = [
            gene,
            f"{result['positive_mean_z']:+.3f}",
            f"{result['negative_mean_z']:+.3f}",
            f"{result['delta_z']:+.3f}",
            fmt_p(float(result["mw_p"])),
            fmt_p(float(result["bh_fdr"])),
        ]
        if result["bh_fdr"] < 0.05:
            interpretation = "Primary-diagnosis subtype signal; not prognosis"
        else:
            interpretation = "Not BH-FDR significant in primary-only analysis"
        values.append(interpretation)
        for index, value in enumerate(values):
            row.cells[index].text = value


def update_summary_tables(document: Document) -> None:
    table1 = document.tables[0]
    for row in table1.rows:
        metric = row.cells[0].text.strip()
        if metric == "Baseline RNA / paired RNA":
            row.cells[1].text = (
                "11 all-relapsed prognostic baseline patients / 8 pairs; "
                "0 FDR-significant adverse-outcome genes"
            )
        if metric == "Primary interpretation":
            row.cells[1].text = (
                "+16 is an external-validation candidate; nested-CV and multiplicity "
                "analyses identify no validated adverse-outcome biomarker"
            )

    table4 = document.tables[3]
    for row in list(table4.rows):
        metric = row.cells[0].text.strip()
        if metric == "Complex karyotype":
            row.cells[1].text = "2 patients; 1 event; HR=11.15; P=0.032; BH-FDR=0.224"
            row.cells[2].text = "Nominal, extremely sparse, and non-significant after multiplicity correction"
        elif metric == "Chromosome 16 gain":
            row.cells[1].text = "Absent in the historical ETV6::RUNX1-positive outcome cohort"
            row.cells[2].text = "Cannot independently validate or refute the public-data candidate"
        elif metric == "del6q":
            row.cells[1].text = "One historical ETV6::RUNX1-positive patient without an event"
            row.cells[2].text = "Too rare for prognostic validation"
        elif metric == "Contemporary registry":
            remove_table_row(table4, row)


def add_table_5(document: Document, summary: dict) -> None:
    survival = summary["survival_ml"]
    expression_ml = summary["expression_ml"]
    anchor = find_paragraph(document, "Supplementary Figures Retained for Transparency")
    heading = document.add_paragraph(
        "Table 5. Multiplicity-controlled and machine-learning reliability assessment",
        style="Heading 2",
    )
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Analysis", "Result", "Interpretation"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    rows = [
        (
            "Adverse-outcome multiple testing",
            "0 BH-FDR-significant families across karyotype, CNA, prognostic RNA, paired RNA, DepMap, and single-center Cox screens",
            "No validated adverse-outcome biomarker",
        ),
        (
            "Unsupervised karyotype clustering",
            "Best k=4; silhouette=0.331; median stability ARI=0.114; event permutation P=1.000; log-rank P=0.942",
            "Weak, unstable clusters without outcome association",
        ),
        (
            "Three-variable clinical-threshold nested-CV Cox",
            f"Mean outer-fold C-index={survival['clinical_mean_c_index']:.3f}",
            "Limited baseline discrimination",
        ),
        (
            "Clinical + chromosome 16 gain",
            f"Mean outer-fold C-index={survival['gain16_mean_c_index']:.3f}; "
            f"paired mean delta={survival['gain16_mean_delta']:+.3f}; "
            f"median delta={survival['gain16_median_delta']:+.3f}; "
            f"outer-fold range {survival['gain16_delta_min']:+.3f} to {survival['gain16_delta_max']:+.3f}",
            "Possible low-frequency signal, but out-of-sample gain is unstable",
        ),
        (
            "Clinical + all eligible karyotype features",
            f"Mean outer-fold C-index={survival['karyotype_mean_c_index']:.3f}; "
            f"paired mean delta={survival['karyotype_mean_delta']:+.3f}; "
            f"median delta={survival['karyotype_median_delta']:+.3f}",
            "No prognostic improvement",
        ),
        (
            "Primary-only subtype expression",
            "11 positive vs 99 negative; SLC7A11 delta Z=+0.663; P=0.000907; BH-FDR=0.0229; BY-FDR=0.244; Holm P=1.000",
            "Subtype-expression candidate; not adverse-outcome or mechanistic validation",
        ),
        (
            "Primary-only supervised RNA classification",
            f"Mean nested-CV outer-fold AUC={expression_ml['nested_cv_auc_mean']:.3f}; "
            f"complete nested-pipeline permutation P={expression_ml['nested_pipeline_permutation_p']:.3g} "
            f"({expression_ml['nested_pipeline_permutations']} permutations)",
            "Sparse fusion-subtype signal; not prognosis",
        ),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    note = document.add_paragraph(
        "ARI: adjusted Rand index; AUC: area under the receiver-operating-characteristic curve; "
        "BH: Benjamini-Hochberg; BY: Benjamini-Yekutieli; CV: cross-validation."
    )
    move_before(anchor, heading._p)
    move_before(anchor, table._tbl)
    move_before(anchor, note._p)


def main() -> None:
    configure_paths(parse_args())
    for required in [SOURCE, FIGURE_5, FIGURE_6, EXPRESSION, SUMMARY, OLD_FIGURE_3, NEW_FIGURE_3]:
        if not required.exists():
            raise FileNotFoundError(required)
    summary = json.loads(SUMMARY.read_text(encoding="utf-8"))
    survival = summary["survival_ml"]
    expression_ml = summary["expression_ml"]
    document = Document(SOURCE)
    replace_embedded_image(document, OLD_FIGURE_3, NEW_FIGURE_3)

    replace_paragraph(
        document,
        "ETV6::RUNX1-positive pediatric B-cell acute lymphoblastic leukemia",
        f"ETV6::RUNX1-positive pediatric B-cell acute lymphoblastic leukemia (B-ALL) usually has a favorable outcome, but a subset of patients experiences relapse. We performed an exploratory reanalysis of public TARGET, NOPHO, DepMap, and cBioPortal data and assessed candidate signals in a de-identified historical single-center ETV6::RUNX1-positive outcome cohort. Reliability was strengthened using Benjamini-Hochberg, Benjamini-Yekutieli, and Holm multiplicity correction, unsupervised stability analysis, and leakage-resistant nested cross-validation. No adverse-outcome candidate from the karyotype, CNA, prognostic RNA, paired RNA, DepMap, or single-center screens survived BH-FDR correction. Chromosome 16 gain (+16) remained the leading low-frequency candidate (HR=4.20, P=0.00879; BH-FDR=0.308), but adding +16 to a three-variable clinical-threshold Cox model produced an unstable mean outer-fold C-index increase of {survival['gain16_mean_delta']:.3f} with median change {survival['gain16_median_delta']:.3f}, whereas adding all eligible karyotype features did not improve discrimination. Unsupervised karyotype clusters were weak, unstable, and unrelated to EFS. In a stricter primary-diagnosis expression analysis using one sample per patient (11 ETV6::RUNX1-positive vs 99 negative), SLC7A11 remained BH-FDR significant but not significant under BY or Holm correction; supervised RNA classification identified a fusion-subtype signal, not a prognostic signal. These integrated findings do not establish a validated adverse-outcome biomarker or therapeutic mechanism and nominate +16 only for multicenter validation.",
    )
    replace_paragraph(
        document,
        "Keywords:",
        "Keywords: ETV6::RUNX1, pediatric B-ALL, multiple testing, nested cross-validation, chromosome 16 gain, reproducibility, prognostic biomarker",
    )
    replace_paragraph(
        document,
        "In this study, we reanalyzed",
        "In this study, we reanalyzed public TARGET, NOPHO, DepMap, and cBioPortal resources and assessed relevant candidates in a de-identified single-center clinical dataset to: (1) screen recurrent karyotype features for association with EFS in ETV6::RUNX1-positive patients; (2) test the robustness and external assessability of del6q, +16, and a previously proposed multi-hit score; (3) screen baseline CNA and RNA data for candidate adverse-outcome biomarkers while controlling false discovery; (4) quantify unsupervised structure and out-of-sample predictive value using stability analysis and nested cross-validation; and (5) distinguish reproducible subtype observations from hypotheses requiring multicenter clinical and experimental validation.",
    )
    replace_paragraph(
        document,
        "RNA-seq analyses.",
        "RNA-seq analyses. Public cBioPortal RPKM Z-scores, rather than raw counts, were analyzed. For prognostic screening, only initial-diagnosis sample types 09 or 03 were retained, with recurrent sample type 04 excluded and one baseline sample selected per patient. Baseline expression was correlated with time to relapse, and paired baseline-relapse samples were assessed using paired Wilcoxon tests. Because all 11 available prognostic baseline RNA patients subsequently relapsed, these analyses cannot compare relapsing with non-relapsing patients. For subtype-expression analysis, a stricter primary-only cohort was derived by retaining one diagnosis sample per patient, prioritizing bone marrow type 09 over peripheral-blood type 03; this yielded 11 ETV6::RUNX1-positive and 99 negative patients. Subtype-expression differences were not interpreted as prognosis or del6q mechanism validation.",
    )
    replace_paragraph(
        document,
        "Statistics.",
        "Statistics. All tests were two-sided. Multiplicity was assessed within biologically and statistically distinct hypothesis families using Benjamini-Hochberg FDR, with Benjamini-Yekutieli FDR and Holm family-wise error correction as conservative sensitivity analyses. Nominal P values are reported only to describe candidate ranking and are not treated as confirmatory when adjusted results are non-significant. Effect estimates are reported with 95% confidence intervals where available. Analyses used Python 3.11 with scipy, lifelines, statsmodels, pandas, numpy, and matplotlib. The original network simulation and treatment framework were retained only as supplementary hypothesis-generating diagrams and were excluded from inferential conclusions.",
    )
    replace_paragraph(
        document,
        "Single-center clinical cohorts.",
        "Single-center historical outcome cohort. The historical outcome cohort comprised 73 TEL/AML1-positive patients diagnosed during 2009-2019 and was used for exploratory EFS and overall-survival analyses. EFS was measured from diagnosis to first relapse or death, with other observations censored at last recorded follow-up. Kaplan-Meier estimates, log-rank tests, Wilson confidence intervals, and univariable Cox models were used. +16 and del6q were prespecified from the public-data analysis. All single-center inferential analyses were considered exploratory because only seven EFS events occurred.",
    )
    ml_anchor = find_paragraph(document, "Single-center historical outcome cohort.")
    ml_anchor.insert_paragraph_before(
        f"Machine-learning reliability analysis. Unsupervised karyotype analysis used Jaccard distance and average-linkage hierarchical clustering after prevalence filtering and removal of duplicate binary feature patterns. Candidate cluster numbers k=2-6 were compared by silhouette, outcome association was evaluated by label permutation and log-rank testing, and stability was quantified by adjusted Rand index across 300 repeated 80% feature-subsampling analyses. Prognostic supervised learning was restricted to the 141-patient TARGET karyotype cohort with 24 EFS events. Three-variable clinical-threshold, clinical-plus-+16, and clinical-plus-all-eligible-karyotype Cox models were evaluated using repeated nested five-fold cross-validation; feature filtering, imputation, and ridge-penalty selection occurred within training folds. Concordance was calculated separately in each outer test fold before aggregation, because partial-hazard scales from independently fitted Cox models are not directly comparable across folds. Primary-only RNA subtype classification used a nearest-centroid classifier with nested feature-number selection. The complete nested classification pipeline was evaluated using {expression_ml['nested_pipeline_permutations']} label permutations, and AUC was calculated separately in each outer test fold before aggregation. The seven-event single-center cohort and 11-patient relapse-only RNA subset were not used to train prognostic machine-learning models.",
        style="Normal",
    )

    replace_paragraph(
        document,
        "A separate subtype-expression comparison reproduced",
        "A stricter primary-only subtype-expression comparison included 11 ETV6::RUNX1-positive and 99 negative patients, with one diagnosis sample retained per patient. It identified 1,346 BH-FDR-significant genes, 485 BY-FDR-significant genes, and 140 Holm-significant genes. SLC7A11 remained higher in ETV6::RUNX1-positive disease (Delta Z=+0.663, P=0.000907, BH-FDR=0.0229), but was not significant under BY-FDR (0.244) or Holm correction (1.000). IGF2BP1 remained significant under all three correction approaches, whereas GCLM was BH- and BY-significant but not Holm-significant. These observations describe a diagnostic fusion-subtype expression program and cannot validate prognosis or a del6q/FOXO3-SLC7A11 mechanism.",
    )
    replace_paragraph(
        document,
        "Within the 18-sample",
        "The outcome-focused RNA dataset remained unsuitable for prognostic machine learning because all 11 eligible baseline patients relapsed. The strict separation between subtype-expression analysis and adverse-outcome analysis was therefore maintained throughout interpretation.",
    )
    replace_paragraph(
        document,
        "Complex karyotype was recorded",
        "Complex karyotype was recorded in two historical patients with one EFS event and yielded a nominal univariable association (HR=11.15, 95%CI 1.23-101.00, P=0.032). After correction across the seven estimable single-center Cox comparisons, BH-FDR was 0.224. Given the extremely sparse positive group, wide confidence interval, and non-significant adjusted result, this finding is hypothesis-generating and is not considered validation.",
    )

    discussion = find_paragraph(document, "Discussion")
    discussion.insert_paragraph_before(
        "Multiplicity-controlled and machine-learning reliability analyses",
        style="Heading 2",
    )
    discussion.insert_paragraph_before(
        "No adverse-outcome screening family yielded a BH-FDR-significant candidate. This included the TARGET karyotype, genome-wide CNA, time-to-relapse RNA, paired relapse RNA, DepMap dependency, and single-center Cox families; BY and Holm sensitivity analyses were at least as conservative. The nominal single-center complex-karyotype association had BH-FDR=0.224, and the +16 karyotype association had BH-FDR=0.308.",
        style="Normal",
    )
    discussion.insert_paragraph_before(
        f"Unsupervised analysis of 36 non-duplicate recurrent karyotype features selected four clusters by silhouette, but separation was weak (silhouette=0.331), feature-subsampling stability was low (median adjusted Rand index=0.114), and clusters were unrelated to EFS events (permutation P=1.000; log-rank P=0.942). Repeated nested cross-validation showed limited three-variable clinical-threshold discrimination (mean outer-fold C-index={survival['clinical_mean_c_index']:.3f}). Adding +16 increased the mean outer-fold C-index to {survival['gain16_mean_c_index']:.3f}, but the paired change had mean {survival['gain16_mean_delta']:+.3f}, median {survival['gain16_median_delta']:+.3f}, and outer-fold range {survival['gain16_delta_min']:+.3f} to {survival['gain16_delta_max']:+.3f}. Adding all eligible karyotype features yielded mean outer-fold C-index={survival['karyotype_mean_c_index']:.3f} and did not improve prediction (Figure 5; Table 5).",
        style="Normal",
    )
    insert_picture_before(document, discussion, FIGURE_5, 6.6)
    discussion.insert_paragraph_before(
        "Figure 5. Multiplicity-controlled and machine-learning reliability assessment. (A) Minimum nominal P values and BH-FDR values across analysis families. (B) Principal-component representation of unsupervised karyotype structure; clusters were weak, unstable, and unrelated to outcome. (C) Primary-only RNA principal components did not show significant global unsupervised separation by fusion status. (D) Repeated nested-CV survival performance for clinical-only, clinical-plus-+16, and clinical-plus-all-karyotype models.",
        style="Normal",
    )
    discussion.insert_paragraph_before(
        f"Primary-only RNA did not show significant global unsupervised separation by fusion status (label-permutation P=0.241), but supervised feature selection identified a sparse fusion-subtype expression signal (mean nested-CV outer-fold AUC={expression_ml['nested_cv_auc_mean']:.3f}; complete nested-pipeline permutation P={expression_ml['nested_pipeline_permutation_p']:.3g}, {expression_ml['nested_pipeline_permutations']} permutations). Primary sample-type distributions were similar by fusion status (Fisher P=0.725). This internal machine-learning result supports subtype biology only and cannot be interpreted as adverse-outcome prediction (Figure 6).",
        style="Normal",
    )
    insert_picture_before(document, discussion, FIGURE_6, 6.6)
    discussion.insert_paragraph_before(
        "Figure 6. Cross-validated supervised-learning results. (A) Repeated nested-CV survival discrimination showed unstable incremental performance from +16 and no improvement from the full eligible karyotype feature set. (B) Primary-only RNA classified ETV6::RUNX1 fusion subtype with high nested-CV AUC, but this analysis did not include an adverse-outcome endpoint.",
        style="Normal",
    )

    replace_paragraph(
        document,
        "This reproducible public-data reanalysis did not identify",
        "This reproducible public-data reanalysis did not identify a validated genomic or expression biomarker of adverse outcome in ETV6::RUNX1-positive pediatric B-ALL. This conclusion was unchanged after applying BH, BY, and Holm multiplicity correction, unsupervised stability analysis, and leakage-resistant nested cross-validation. The strongest karyotype candidate was chromosome 16 gain, which showed an adverse-risk direction across primary and leave-one-positive-patient-out analyses but occurred in only seven patients, did not survive screening-wide FDR correction, and did not provide a consistently positive out-of-sample performance gain. The appropriate conclusion is nomination for external validation, not clinical validation.",
    )
    replace_paragraph(
        document,
        "SLC7A11 upregulation is",
        "SLC7A11 upregulation remains a reproducible but correction-sensitive ETV6::RUNX1 subtype-expression observation. Under the stricter primary-only analysis it was significant by BH-FDR but not by BY or Holm correction. It does not validate a del6q/FOXO3 mechanism because the available expression and outcome cohorts do not provide the required concurrent lesion, expression, and independent clinical validation. Cross-version DepMap analyses likewise did not reproduce SLC7A11 or LATS1 as del6q-specific dependencies after multiple-testing correction.",
    )
    replace_paragraph(
        document,
        "Additional limitations include",
        "Additional limitations include automated rather than blinded manual ISCN review, low numbers of patients with specific karyotype lesions, only 24 EFS events in the TARGET karyotype cohort, gene-level rather than uniformly reprocessed segment-level TARGET CNA calls, unavailable DepMap 24Q2 bulk files, changing model availability across DepMap releases, and use of public RPKM Z-scores rather than raw RNA counts. Although nested cross-validation reduced model-selection optimism, the supervised analyses remain internal validation and do not replace an independent cohort. The historical single-center outcome cohort was retrospective, contained heterogeneous source-recorded karyotype and MRD fields, had only seven EFS events, and lacked RNA-seq for expression-biomarker validation.",
    )
    replace_paragraph(
        document,
        "The immediate next step is",
        "The immediate next step is a prespecified multicenter external validation study containing complete karyotype, segment-level CNA, Day-29 MRD, treatment variables, and diagnosis RNA-seq from both relapsing and non-relapsing ETV6::RUNX1-positive patients. Chromosome 16 gain should be prioritized for blinded cytogenetic review, but model performance must be assessed in an untouched external cohort. Functional studies of FOXO3, SLC7A11, and ferroptosis remain biologically plausible research directions, but the current data do not justify changes to FISH panels, risk assignment, or targeted therapy.",
    )
    replace_paragraph(
        document,
        "Data sharing:",
        "Data availability. Public source data are available from the cBioPortal DataHub TARGET Phase II study, the NOPHO/Oksa Zenodo records, and official DepMap Public releases, with source versions and commit identifiers listed in the reproducibility repository at https://github.com/ABLE-Du/ETV6-RUNX1-biomarker-reanalysis. Reanalysis scripts use explicit source-root and workspace-root parameters rather than machine-specific paths. The repository includes a machine-readable input-file checksum manifest, software-environment record, public-data-derived result tables, aggregate-only single-center results, and figure-generation code. De-identified individual-level single-center clinical data are not publicly released because of privacy and institutional-governance requirements.",
    )

    replace_paragraph(
        document,
        "Table 3. Reproducible ETV6::RUNX1 subtype-expression differences",
        "Table 3. Primary-only ETV6::RUNX1 subtype-expression differences in public TARGET RNA-seq Z-scores (n=11 positive vs. 99 negative)",
    )
    replace_paragraph(
        document,
        "Delta Z is the difference",
        "Delta Z is the difference in mean public cBioPortal RPKM Z-score using one primary-diagnosis sample per patient. BH-FDR is genome-wide across 23,590 unique expression genes. BY-FDR and Holm correction are reported in Table 5 and the supplementary analysis outputs. These results describe subtype expression and do not validate prognosis or a del6q mechanism.",
    )
    update_expression_table(document)
    update_summary_tables(document)
    add_table_5(document, summary)

    replace_all_paragraphs(
        document,
        "Figure 3. Single-center outcomes in ETV6::RUNX1-positive pediatric B-ALL.",
        "Figure 3. Historical single-center outcomes in ETV6::RUNX1-positive pediatric B-ALL. (a) Historical outcome-cohort composition by treatment era. (b) Overall EFS and overall survival among 73 historical patients. (c) Exploratory EFS curves by source-recorded risk group. (d) Patient-level timing and type of the seven EFS events.",
    )
    replace_all_paragraphs(
        document,
        "Supplementary Figure S6. Contemporary single-center registry phenotype and early response.",
        "Supplementary Figure S6. Non-ETV6::RUNX1-selected contemporary B-ALL registry phenotype and early response. (a) Diagnostic immunophenotype heatmap among 42 contemporary B-ALL registrations; blank cells indicate unavailable measurements. (b) Marker prevalence using a threshold of at least 20% positive cells, with denominators restricted to patients with available measurements. (c) D19 and D46 MRD distributions. This background registry was not selected by ETV6::RUNX1 status, was excluded from the main analysis and primary tables, and does not provide outcome or biomarker validation.",
    )

    legends_anchor = find_paragraph(document, "Tables")
    legends_anchor.insert_paragraph_before(
        "Figure 5. Multiplicity-controlled and machine-learning reliability assessment. (A) Minimum nominal P values and BH-FDR values across analysis families. (B) Principal-component representation of unsupervised karyotype structure; clusters were weak, unstable, and unrelated to outcome. (C) Primary-only RNA principal components did not show significant global unsupervised separation by fusion status. (D) Repeated nested-CV survival performance for clinical-only, clinical-plus-+16, and clinical-plus-all-karyotype models.",
        style="Normal",
    )
    legends_anchor.insert_paragraph_before(
        "Figure 6. Cross-validated supervised-learning results. (A) Repeated nested-CV survival discrimination showed unstable incremental performance from +16 and no improvement from the full eligible karyotype feature set. (B) Primary-only RNA classified ETV6::RUNX1 fusion subtype with high nested-CV AUC, but this analysis did not include an adverse-outcome endpoint.",
        style="Normal",
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
